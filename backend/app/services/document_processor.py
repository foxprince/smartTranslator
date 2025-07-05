"""
文档处理服务
支持多种文档格式的上传、解析和翻译
"""
class DocumentProcessingError(Exception):
    """文档处理异常"""
    pass


import os
import uuid
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import logging
from datetime import datetime

# 文档处理库
import PyPDF2
import docx
import openpyxl
import pandas as pd
from bs4 import BeautifulSoup
import markdown
import json
import xml.etree.ElementTree as ET

# 图像处理和OCR
from PIL import Image
import pytesseract

# 异步支持
import asyncio
from concurrent.futures import ThreadPoolExecutor

from app.core.config import settings
from app.services.translation_engine import TranslationEngine
from app.schemas.translation import TranslationRequest, LanguageCode, TranslationProvider

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器"""
    
    # 支持的文档格式
    SUPPORTED_FORMATS = {
        # 文本文档
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.rtf': 'application/rtf',
        
        # Office文档
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xls': 'application/vnd.ms-excel',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.ppt': 'application/vnd.ms-powerpoint',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        
        # PDF文档
        '.pdf': 'application/pdf',
        
        # 网页文档
        '.html': 'text/html',
        '.htm': 'text/html',
        '.xml': 'application/xml',
        
        # 数据格式
        '.json': 'application/json',
        '.csv': 'text/csv',
        
        # 图像文档（OCR）
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.bmp': 'image/bmp',
        '.tiff': 'image/tiff',
        '.gif': 'image/gif',
    }
    
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR or "uploads")
        self.upload_dir.mkdir(exist_ok=True)
        
        self.processed_dir = self.upload_dir / "processed"
        self.processed_dir.mkdir(exist_ok=True)
        
        self.translation_engine = TranslationEngine()
        self.executor = ThreadPoolExecutor(max_workers=4)
    
    async def upload_document(
        self, 
        file_content: bytes, 
        filename: str,
        user_id: str,
        project_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        上传文档
        
        Args:
            file_content: 文件内容
            filename: 文件名
            user_id: 用户ID
            project_id: 项目ID（可选）
            
        Returns:
            上传结果信息
        """
        try:
            # 验证文件格式
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.SUPPORTED_FORMATS:
                raise ValueError(f"不支持的文件格式: {file_ext}")
            
            # 验证文件大小
            max_size = getattr(settings, 'MAX_UPLOAD_SIZE', 50 * 1024 * 1024)  # 50MB
            if len(file_content) > max_size:
                raise ValueError(f"文件大小超过限制: {len(file_content)} > {max_size}")
            
            # 生成唯一文件ID
            file_id = str(uuid.uuid4())
            
            # 保存原始文件
            original_path = self.upload_dir / f"{file_id}_{filename}"
            with open(original_path, 'wb') as f:
                f.write(file_content)
            
            # 创建文档记录
            document_info = {
                "id": file_id,
                "original_filename": filename,
                "file_path": str(original_path),
                "file_size": len(file_content),
                "file_type": file_ext,
                "mime_type": self.SUPPORTED_FORMATS[file_ext],
                "user_id": user_id,
                "project_id": project_id,
                "upload_time": datetime.now().isoformat(),
                "status": "uploaded",
                "processing_status": "pending"
            }
            
            logger.info(f"文档上传成功: {filename} ({file_id})")
            
            return document_info
            
        except Exception as e:
            logger.error(f"文档上传失败: {e}")
            raise
    
    async def extract_text_from_document(self, document_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        从文档中提取文本内容
        
        Args:
            document_info: 文档信息
            
        Returns:
            提取结果
        """
        try:
            file_path = Path(document_info["file_path"])
            file_type = document_info["file_type"]
            
            logger.info(f"开始提取文本: {file_path}")
            
            # 根据文件类型选择提取方法
            if file_type == '.txt':
                extracted_text = await self._extract_from_txt(file_path)
            elif file_type == '.pdf':
                extracted_text = await self._extract_from_pdf(file_path)
            elif file_type in ['.doc', '.docx']:
                extracted_text = await self._extract_from_word(file_path)
            elif file_type in ['.xls', '.xlsx']:
                extracted_text = await self._extract_from_excel(file_path)
            elif file_type in ['.html', '.htm']:
                extracted_text = await self._extract_from_html(file_path)
            elif file_type == '.md':
                extracted_text = await self._extract_from_markdown(file_path)
            elif file_type == '.json':
                extracted_text = await self._extract_from_json(file_path)
            elif file_type == '.csv':
                extracted_text = await self._extract_from_csv(file_path)
            elif file_type == '.xml':
                extracted_text = await self._extract_from_xml(file_path)
            elif file_type in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                extracted_text = await self._extract_from_image(file_path)
            else:
                raise ValueError(f"不支持的文件类型: {file_type}")
            
            # 更新文档信息
            result = {
                **document_info,
                "extracted_text": extracted_text,
                "text_length": len(extracted_text),
                "processing_status": "text_extracted",
                "extraction_time": datetime.now().isoformat()
            }
            
            logger.info(f"文本提取完成: {len(extracted_text)} 字符")
            
            return result
            
        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            raise
    
    async def _extract_from_txt(self, file_path: Path) -> str:
        """从TXT文件提取文本"""
        def extract():
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def _extract_from_pdf(self, file_path: Path) -> str:
        """从PDF文件提取文本"""
        def extract():
            text = ""
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e:
                logger.warning(f"PDF提取失败，尝试OCR: {e}")
                # 如果PDF提取失败，尝试转换为图像后OCR
                # 这里可以添加PDF转图像的逻辑
                pass
            return text.strip()
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def _extract_from_word(self, file_path: Path) -> str:
        """从Word文档提取文本"""
        def extract():
            try:
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\t"
                        text += "\n"
                
                return text.strip()
            except Exception as e:
                logger.error(f"Word文档提取失败: {e}")
                return ""
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def _extract_from_excel(self, file_path: Path) -> str:
        """从Excel文件提取文本"""
        def extract():
            try:
                workbook = openpyxl.load_workbook(file_path)
                text = ""
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"=== {sheet_name} ===\n"
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip():
                            text += row_text + "\n"
                    text += "\n"
                
                return text.strip()
            except Exception as e:
                logger.error(f"Excel文件提取失败: {e}")
                return ""
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def _extract_from_html(self, file_path: Path) -> str:
        """从HTML文件提取文本"""
        def extract():
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    html_content = f.read()
                
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # 移除脚本和样式标签
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # 提取文本
                text = soup.get_text()
                
                # 清理文本
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text
            except Exception as e:
                logger.error(f"HTML文件提取失败: {e}")
                return ""
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def _extract_from_markdown(self, file_path: Path) -> str:
        """从Markdown文件提取文本"""
        def extract():
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    md_content = f.read()
                
                # 转换为HTML然后提取文本
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
            except Exception as e:
                logger.error(f"Markdown文件提取失败: {e}")
                return ""
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def _extract_from_json(self, file_path: Path) -> str:
        """从JSON文件提取文本"""
        def extract():
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    data = json.load(f)
                
                def extract_text_from_json(obj, path=""):
                    texts = []
                    if isinstance(obj, dict):
                        for key, value in obj.items():
                            new_path = f"{path}.{key}" if path else key
                            texts.extend(extract_text_from_json(value, new_path))
                    elif isinstance(obj, list):
                        for i, item in enumerate(obj):
                            new_path = f"{path}[{i}]"
                            texts.extend(extract_text_from_json(item, new_path))
                    elif isinstance(obj, str):
                        texts.append(f"{path}: {obj}")
                    else:
                        texts.append(f"{path}: {str(obj)}")
                    return texts
                
                text_lines = extract_text_from_json(data)
                return "\n".join(text_lines)
            except Exception as e:
                logger.error(f"JSON文件提取失败: {e}")
                return ""
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def _extract_from_csv(self, file_path: Path) -> str:
        """从CSV文件提取文本"""
        def extract():
            try:
                df = pd.read_csv(file_path, encoding='utf-8', errors='ignore')
                
                # 将DataFrame转换为文本
                text = ""
                
                # 添加列标题
                text += "\t".join(df.columns) + "\n"
                
                # 添加数据行
                for _, row in df.iterrows():
                    row_text = "\t".join([str(val) if pd.notna(val) else "" for val in row])
                    text += row_text + "\n"
                
                return text.strip()
            except Exception as e:
                logger.error(f"CSV文件提取失败: {e}")
                return ""
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def _extract_from_xml(self, file_path: Path) -> str:
        """从XML文件提取文本"""
        def extract():
            try:
                tree = ET.parse(file_path)
                root = tree.getroot()
                
                def extract_text_from_element(element, path=""):
                    texts = []
                    current_path = f"{path}/{element.tag}" if path else element.tag
                    
                    if element.text and element.text.strip():
                        texts.append(f"{current_path}: {element.text.strip()}")
                    
                    for child in element:
                        texts.extend(extract_text_from_element(child, current_path))
                    
                    return texts
                
                text_lines = extract_text_from_element(root)
                return "\n".join(text_lines)
            except Exception as e:
                logger.error(f"XML文件提取失败: {e}")
                return ""
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def _extract_from_image(self, file_path: Path) -> str:
        """从图像文件提取文本（OCR）"""
        def extract():
            try:
                # 使用Tesseract OCR提取文本
                image = Image.open(file_path)
                
                # 图像预处理（可选）
                # image = image.convert('L')  # 转换为灰度
                
                # OCR识别
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                
                return text.strip()
            except Exception as e:
                logger.error(f"图像OCR提取失败: {e}")
                return ""
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract)
    
    async def translate_document(
        self,
        document_info: Dict[str, Any],
        source_language: LanguageCode,
        target_language: LanguageCode,
        provider: TranslationProvider = TranslationProvider.GOOGLE,
        chunk_size: int = 1000
    ) -> Dict[str, Any]:
        """
        翻译文档内容
        
        Args:
            document_info: 文档信息
            source_language: 源语言
            target_language: 目标语言
            provider: 翻译提供商
            chunk_size: 文本分块大小
            
        Returns:
            翻译结果
        """
        try:
            extracted_text = document_info.get("extracted_text", "")
            if not extracted_text:
                raise ValueError("文档中没有可翻译的文本内容")
            
            # 将长文本分块处理
            text_chunks = self._split_text_into_chunks(extracted_text, chunk_size)
            
            logger.info(f"开始翻译文档，共 {len(text_chunks)} 个文本块")
            
            # 创建翻译请求
            translation_request = TranslationRequest(
                texts=text_chunks,
                source_language=source_language,
                target_language=target_language,
                provider=provider
            )
            
            # 执行翻译
            translation_result = await self.translation_engine.translate_batch(translation_request)
            
            # 合并翻译结果
            translated_text = ""
            for translation in translation_result.translations:
                translated_text += translation.translated_text + "\n"
            
            # 保存翻译结果
            translated_file_path = self._save_translated_document(
                document_info, 
                translated_text, 
                source_language, 
                target_language
            )
            
            # 更新文档信息
            result = {
                **document_info,
                "translation_result": {
                    "source_language": source_language.value,
                    "target_language": target_language.value,
                    "provider": provider.value,
                    "translated_text": translated_text,
                    "translated_file_path": translated_file_path,
                    "translation_stats": {
                        "total_chunks": len(text_chunks),
                        "success_count": translation_result.success_count,
                        "failed_count": translation_result.failed_count,
                        "total_cost": translation_result.total_cost,
                        "quality_summary": translation_result.quality_summary,
                        "processing_time": translation_result.processing_time
                    }
                },
                "processing_status": "translated",
                "translation_time": datetime.now().isoformat()
            }
            
            logger.info(f"文档翻译完成: {translation_result.success_count}/{len(text_chunks)} 成功")
            
            return result
            
        except Exception as e:
            logger.error(f"文档翻译失败: {e}")
            raise
    
    def _split_text_into_chunks(self, text: str, chunk_size: int) -> List[str]:
        """将长文本分割成小块"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        sentences = text.split('\n')
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += sentence + "\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + "\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _save_translated_document(
        self,
        document_info: Dict[str, Any],
        translated_text: str,
        source_language: LanguageCode,
        target_language: LanguageCode
    ) -> str:
        """保存翻译后的文档"""
        try:
            # 生成翻译文件名
            original_filename = document_info["original_filename"]
            file_stem = Path(original_filename).stem
            file_ext = Path(original_filename).suffix
            
            translated_filename = f"{file_stem}_{source_language.value}_to_{target_language.value}{file_ext}"
            translated_path = self.processed_dir / translated_filename
            
            # 根据原文件格式保存翻译结果
            if file_ext.lower() == '.txt':
                with open(translated_path, 'w', encoding='utf-8') as f:
                    f.write(translated_text)
            elif file_ext.lower() == '.html':
                # 保持HTML格式
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Translated Document</title>
</head>
<body>
    <pre>{translated_text}</pre>
</body>
</html>"""
                with open(translated_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            else:
                # 默认保存为文本文件
                with open(translated_path.with_suffix('.txt'), 'w', encoding='utf-8') as f:
                    f.write(translated_text)
                translated_path = translated_path.with_suffix('.txt')
            
            return str(translated_path)
            
        except Exception as e:
            logger.error(f"保存翻译文档失败: {e}")
            raise
    
    async def get_document_info(self, document_id: str) -> Optional[Dict[str, Any]]:
        """获取文档信息"""
        # 这里应该从数据库获取文档信息
        # 暂时返回None，实际实现需要数据库支持
        return None
    
    async def list_documents(
        self, 
        user_id: str, 
        project_id: Optional[str] = None,
        status: Optional[str] = None,
        limit: int = 50,
        offset: int = 0
    ) -> List[Dict[str, Any]]:
        """列出用户的文档"""
        # 这里应该从数据库查询文档列表
        # 暂时返回空列表，实际实现需要数据库支持
        return []
    
    async def delete_document(self, document_id: str, user_id: str) -> bool:
        """删除文档"""
        try:
            # 获取文档信息
            document_info = await self.get_document_info(document_id)
            if not document_info:
                return False
            
            # 检查权限
            if document_info["user_id"] != user_id:
                raise PermissionError("无权限删除此文档")
            
            # 删除文件
            file_path = Path(document_info["file_path"])
            if file_path.exists():
                file_path.unlink()
            
            # 删除翻译文件（如果存在）
            if "translation_result" in document_info:
                translated_file_path = document_info["translation_result"].get("translated_file_path")
                if translated_file_path:
                    translated_path = Path(translated_file_path)
                    if translated_path.exists():
                        translated_path.unlink()
            
            # 从数据库删除记录
            # 这里需要数据库操作
            
            logger.info(f"文档删除成功: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"删除文档失败: {e}")
            return False
    
    def get_supported_formats(self) -> Dict[str, str]:
        """获取支持的文件格式"""
        return self.SUPPORTED_FORMATS.copy()
    
    async def get_processing_stats(self) -> Dict[str, Any]:
        """获取处理统计信息"""
        try:
            # 统计上传目录中的文件
            total_files = len(list(self.upload_dir.glob("*")))
            processed_files = len(list(self.processed_dir.glob("*")))
            
            # 计算总文件大小
            total_size = sum(f.stat().st_size for f in self.upload_dir.glob("*") if f.is_file())
            processed_size = sum(f.stat().st_size for f in self.processed_dir.glob("*") if f.is_file())
            
            return {
                "total_files": total_files,
                "processed_files": processed_files,
                "total_size_bytes": total_size,
                "processed_size_bytes": processed_size,
                "supported_formats": len(self.SUPPORTED_FORMATS),
                "upload_directory": str(self.upload_dir),
                "processed_directory": str(self.processed_dir)
            }
            
        except Exception as e:
            logger.error(f"获取处理统计失败: {e}")
            return {}


# 全局文档处理器实例
document_processor = DocumentProcessor()
